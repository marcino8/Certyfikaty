import openpyxl
from pathlib import Path
from docx import Document
import os
import comtypes.client
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

emaile = []
imiona = []
nazwiska = []
wzor = ''


def wczytajDane():
    xlsx_file = Path('odp.xlsx')
    wb_obj = openpyxl.load_workbook(xlsx_file)
    sheet = wb_obj.active
    for row in sheet.iter_rows(max_row=sheet.max_row):
        if row[4].value == "Tak":
            emaile.append(row[1].value)
            imiona.append(row[2].value)
            nazwiska.append(row[3].value)


def generujKorespondencje():

    for i in range(len(imiona)):
        document = Document('zaswiadczenie.docx')
        for p in document.paragraphs:
            if 'XYZ' in p.text:
                p.text = ''
                p.add_run(imiona[i]+" "+nazwiska[i]).bold = True
                for r in p.runs:
                    r.font.name = "Source Sans Pro Light"
                    r.font.size = 228600

        for t in document.tables:
            wzor = t.cell(0, 0).text
            t.cell(0, 0).text = ''
            t.cell(0, 0).add_paragraph().add_run(wzor.replace("1", str(i+1))).bold = False
            for r in t.cell(0, 0).paragraphs[1].runs:
                r.font.name = "Source Sans Pro Light"
                r.font.size = 127000
        document.save("WordCert/new"+str(i+1)+".docx")


def konwertujDoPdf(schemat):
    print(len(imiona))
    for i in range(len(imiona)):
        in_file = os.path.abspath("WordCert/new"+str(i+1)+".docx")
        out_file = os.path.abspath("PdfCert/"+str(i+1)+schemat+".pdf")
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_file)
        doc.SaveAs(out_file, FileFormat=17)
        doc.Close()
        word.Quit()
        print(float(i+1)/float(len(imiona)))

def sendMail(schemat, password):
    mail_content ='''Szanowni Państwo,
    W załączniku znajduje się certyfikat potwierdzający uczestnictwo w wieczorku dendrologicznym.

    Z poważaniem
    Arboris TEAM.

    Mail został wygenerowany automatycznie, w razie pytań prosimy o kontakt na:
    arboris@gmail.com.
    '''
    sender_address = 'rertospy3333@gmail.com'
    sender_pass = password

    for i in range(len(imiona)):
        message = MIMEMultipart()
        receiver_address = emaile[i]
        message['From'] = sender_address
        message['To'] = receiver_address
        message['Subject'] = 'Certyfikat uczestnictwa w wieczorku dendrologicznym'
        message.attach(MIMEText(mail_content, 'plain'))
        attach_file_name = 'PdfCert/'+str(i+1)+schemat+'.pdf'
        attach_file = open(attach_file_name, 'rb')
        payload = MIMEBase('application', 'octate-stream')
        payload.set_payload((attach_file).read())
        encoders.encode_base64(payload)
        payload.add_header('Content-Disposition', 'attachment', filename="Certyfikat.pdf")
        message.attach(payload)
        session = smtplib.SMTP('smtp.gmail.com', 587)  # use gmail with port
        session.starttls()  # enable security
        session.login(sender_address, sender_pass)  # login with mail_id and password
        text = message.as_string()
        session.sendmail(sender_address, receiver_address, text)
        session.quit()


if __name__ == '__main__':
    print("Podaj hasło do maila arboris: ")
    haslo = input()
    print("Podaj schemat uzyty we wzorze zaswiadczenia:")
    schemat = input()
    wczytajDane()
    generujKorespondencje()
    konwertujDoPdf(schemat)
    print("Czy wegenerowane certyfikaty są ok? Konntynuowac? [T/N]")
    kontyn = input()
    if kontyn == "T":
        sendMail(schemat, haslo)
        print("Wysłano maile")


